import pandas as pd
import re
import datetime
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import RGBColor



class Delta:
    """
    This class is created to find differrences and problematic rows between the two main databases of the project. 
    its purpose is to soften out these differences eventually. It can also produce a table with similar titels 
    between the databases which might be the same document. 
    """

    def __init__(self):
        self.relatics = pd.read_excel("needed_files\MDR.xlsx")
        self.sharepoint = pd.read_excel("needed_files\Document Data.xlsx")

    @staticmethod
    def preprocess_string(s):
        """
        Method to clean the data of unwanted characters for the comparison
        """
        # Check if the input is a valid string, otherwise return an empty string
        if isinstance(s, str):
            s = re.sub(r'^E\d{2}\s*-\s*', '', s)
            s = re.sub(r'[^a-zA-Z0-9]', '', s).lower()
            return s
        else:
            return ''

    def delta(self, preprocess_function):
        """
        The delta between the MDR in Sharpoint and Relatics
        :param preprocess_function:
        """


        self.relatics["Preprocessed Title"] = self.relatics["Title"].apply(preprocess_function)
        self.sharepoint["Preprocessed Title"] = self.sharepoint["Document Title"].apply(preprocess_function)

        # Preprocess phases in both DataFrames
        self.relatics["Preprocessed Phase"] = self.relatics["Phase"].apply(preprocess_function)
        self.sharepoint["Preprocessed Phase"] = self.sharepoint["Purpose of Issue"].apply(preprocess_function)

        # lines for discipline

        sharepoint_filtered = self.sharepoint
        mismatches_share_title = sharepoint_filtered[
            ~sharepoint_filtered["Preprocessed Title"].isin(self.relatics["Preprocessed Title"])]
        mismatches_share_phase = sharepoint_filtered[
            ~sharepoint_filtered["Preprocessed Phase"].isin(self.relatics["Preprocessed Phase"])]

        mismatches_share_title["Mismatch Type"] = "Title"
        mismatches_share_phase["Mismatch Type"] = "Phase"

        mismatches_share = pd.concat([mismatches_share_title, mismatches_share_phase])

        mismatches_share = mismatches_share[
            ["TenneT Document ID", "Document Title", "Contractor's Discipline", "Purpose of Issue",
             "Mismatch Type"]]

        relatics_filtered = self.relatics
        mismatches_relatics_title = relatics_filtered[
            ~relatics_filtered["Preprocessed Title"].isin(self.sharepoint["Preprocessed Title"])]
        mismatches_relatics_phase = relatics_filtered[
            ~relatics_filtered["Preprocessed Phase"].isin(self.sharepoint["Preprocessed Phase"])]

        mismatches_relatics_title["Mismatch Type"] = "Title"
        mismatches_relatics_phase["Mismatch Type"] = "Phase"

        mismatches_relatics = pd.concat([mismatches_relatics_title, mismatches_relatics_phase])

        mismatches_relatics = mismatches_relatics[
            ["Project Document Number", "Title", "Discipline", "Phase", "Mismatch Type"]]

        self.delta_df_share = pd.DataFrame(mismatches_share)
        self.delta_df_relatics = pd.DataFrame(mismatches_relatics)

        return self.delta_df_share, self.delta_df_relatics

    def share_prob(self):
        """
        This function prints out the problematic rows of the Sharepoint MDR
        """
        sharepoint_analysis = self.sharepoint
        title = sharepoint_analysis[sharepoint_analysis["Document Title"].isna()]
        phase = sharepoint_analysis[sharepoint_analysis["Purpose of Issue"].isna()]
        discipline = sharepoint_analysis[sharepoint_analysis["Contractor's Discipline"].isna()]

        leeg_analysis = pd.concat([title, phase, discipline])
        leeg_analysis = leeg_analysis[
            ["TenneT Document ID", "Document Title", "Purpose of Issue", "Contractor's Discipline"]]
        leeg_analysis["TenneT Document ID"].drop_duplicates(keep='first')
        self.analysis_share_df = pd.DataFrame(leeg_analysis)

        return self.analysis_share_df

    def semantics(self):
        """
        This method conducts a semantics analysis between the titles produced by the delta function
        It produces a list with titles in the two databases which might refer to the same document.
        """
        share_titles_cleaned = [title if isinstance(title, str) else '' for title in
                                self.delta_df_share["Document Title"].tolist()]
        relatics_titles_cleaned = [title if isinstance(title, str) else '' for title in
                                   self.delta_df_relatics["Title"].tolist()]

        # Create a TF-IDF vectorizer
        vectorizer = TfidfVectorizer()

        # Fit and transform document titles
        share_title_vectors = vectorizer.fit_transform(share_titles_cleaned)
        relatics_title_vectors = vectorizer.transform(relatics_titles_cleaned)

        # Calculate cosine similarities
        similarity_matrix = cosine_similarity(share_title_vectors, relatics_title_vectors)

        # Create a DataFrame with columns: Share Title, Relatics Title, Similarity Score
        similarity_data = []
        for i, share_title in enumerate(self.delta_df_share["Document Title"].tolist()):
            for j, relatics_title in enumerate(self.delta_df_relatics["Title"].tolist()):
                similarity_data.append([share_title, relatics_title, similarity_matrix[i][j]])

        similarity_df = pd.DataFrame(similarity_data, columns=["Share Title", "Relatics Title", "Similarity Score"])

        self.filtered_df = similarity_df[
            (similarity_df["Similarity Score"] > 0.8) & (similarity_df["Similarity Score"] < 1)]
        self.filtered_df = self.filtered_df.drop_duplicates()

        return self.filtered_df

    def count_dif(self):
        """
        This method counts the amount of documents found to have diferences in title and phase
        from sharepoint to relatics and the other way around. Or it counts the number of rows of the Sharepoint MDR missing at least
        one the the following: Title, Phase , Discipline. It is used to print out a small on-the-go report foe the user.
        """
        dif_share_rel = len(self.delta_df_share)
        dif_rel_share = len(self.delta_df_relatics)
        flag_title_share = len(self.delta_df_share[self.delta_df_share["Mismatch Type"] == "Title"])
        flag_phase_share = len(self.delta_df_share[self.delta_df_share["Mismatch Type"] == "Phase"])
        flag_title_rel = len(self.delta_df_relatics[self.delta_df_relatics["Mismatch Type"] == "Title"])
        flag_phase_rel = len(self.delta_df_relatics[self.delta_df_relatics["Mismatch Type"] == "Phase"])
        self.totaal = f"There are {dif_share_rel} documents found in Sharepoint but not in Relatics. Out of them {flag_title_share} are flagged on title and {flag_phase_share} on phase.\
There are {dif_rel_share} documents found in Relatics but not Sharepoint. Out of them {flag_title_rel} are flagged on title and {flag_phase_rel} on phase."
        leeg = len(self.analysis_share_df)
        self.rows_leg = f"There are {leeg} rows without Title, Phase or Discipline in the Sharepoint MDR that should be fixed"
        title_similar = len(self.filtered_df)
        self.similar = f"There are {title_similar} similar titles between the two databases"

        return f"{self.rows_leg}. {self.totaal} {self.similar}"

    def download(self, selected_directory):
        """
        Download function taking into acount if a comparison is being made between the files
        or if only the Sharepoint MDR is analysed

        Selected_directory: Used in order for the user to specify the location of the exported file.
        """

        file_path = f"{selected_directory}/MDR Data Analysis_{datetime.date.today()}.xlsx"
        with pd.ExcelWriter(file_path) as writer:
            self.delta_df_share.to_excel(writer, sheet_name="In SP not in Rel", index=True)
            self.delta_df_relatics.to_excel(writer, sheet_name="In Rel not in SP", index=True)
            self.analysis_share_df.to_excel(writer, sheet_name=f"Please fix", index=True)
            self.filtered_df.to_excel(writer, sheet_name=f"Similar titles", index=True)

        file_path1 = f"{selected_directory}/MDR Problems Report {datetime.date.today()}.txt"
        with open(file_path1, "w") as txt_file:
            txt_file.write(f"{self.rows_leg}. {self.totaal} {self.similar}")


class CoverSheet():
    """
    This class is used in order to produce the TenneT coversheet for issued documents. It used the latest version of the 
    Sharepoint MDR as its base. Fuethermore the Coversheet Template is provided by DCC
    """
    def __init__(self):
        self.df = pd.read_excel("needed_files\Document Data.xlsx", sheet_name = "Project documents", engine = 'openpyxl')
        self.doc = Document("needed_files\Coversheet.docx")
        self.export_code_originator = None
        self.export_code_tennet = None

    def title_search(self, user_title):
        """
        This method is used because not all documents hase an originator code in the MDR. It checks which code can the user insert to 
        produce the coversheet based on titles.

        params: 
        user_title: the Title of the documnet the user wishes to issue based on the latest MDR version.
        """
        def get_cell_value_codes(index , name):
            value = self.df.loc[index, name]
            return "N/A" if pd.isna(value) else str(value)
        for i, row in self.df.iterrows():
            #if it only originators: return (please use the following code in the originator's code app)
            if user_title == row["Document Title"]:
                if get_cell_value_codes(i,"Originator's Document ID (may further modifications required)")!= "N/A" and get_cell_value_codes(i,"TenneT Document ID")== "N/A":
                    code = get_cell_value_codes(i,"Originator's Document ID (may further modifications required)")
                    string = f"For {user_title}, please use the line for the Originator's ID along with code: {code}"
                    return string
                #if only tennet no originators
                elif get_cell_value_codes(i,"Originator's Document ID (may further modifications required)") == "N/A" and get_cell_value_codes(i,"TenneT Document ID") != "N/A":
                    code = get_cell_value_codes(i,"TenneT Document ID")
                    string = f"For {user_title}, please use the line for the TenneT Document ID along with code: {code}"
                    return string
                #title has both codes
                elif get_cell_value_codes(i,"Originator's Document ID (may further modifications required)") != "N/A" and get_cell_value_codes(i,"TenneT Document ID") != "N/A":
                    core_or = get_cell_value_codes(i,"Originator's Document ID (may further modifications required)")
                    tennet_cod = get_cell_value_codes(i,"TenneT Document ID")
                    string = f"For {user_title}, please use one of the following: \n{core_or} (in the Originator's Code ID tab) \n{tennet_cod} (in the TenneT Document Code tab)"
                    return string
        else:
            string = "Not a valid Title"
            return string

    def doc_title_originator(self, originator_title):
        df = pd.read_excel("needed_files\Document Data.xlsx", sheet_name = "Project documents", engine = 'openpyxl')

        def get_cell_value(index , name):
            value = df.loc[index, name]
            return "N/A" if pd.isna(value) else str(value)
        
        for i, row in df.iterrows():
            if originator_title == row["Originator's Document ID (may further modifications required)"]:
                cell_value1 = get_cell_value(i, "Document Title")
                break
            elif originator_title not in str(row["Originator's Document ID (may further modifications required)"]):
                cell_value1 = "Not a valid Originato's Code" 

        return cell_value1
    
    def doc_title_tennet(self, tennet_title):
        df = pd.read_excel("needed_files\Document Data.xlsx", sheet_name = "Project documents", engine = 'openpyxl')

        def get_cell_value(index , name):
            value = df.loc[index, name]
            return "N/A" if pd.isna(value) else str(value)

        for i, row in df.iterrows():
            if tennet_title == row["Originator's Document ID (may further modifications required)"]:
                cell_value1 = get_cell_value(i, "Document Title")
                break
            elif tennet_title not in str(row["Originator's Document ID (may further modifications required)"]):
                cell_value1 = "Not a valid Originato's Code" 

        return cell_value1
    
    
    def table_originator(self, originator):
        """
        This method is used to construct the main table of the coversheet for documents with only an Originato's title

        params: 
        originator: The originator's code based on the document's titles
        """
        # Get the table and cell to populate
        table = self.doc.tables[0]
        cell_to_populate1 = table.cell(0, 1)
        cell_to_populate2 = table.cell(1, 0)
        cell_to_populate3 = table.cell(1, 1)

        cell_to_populate4 = table.cell(2, 0)
        cell_to_populate5 = table.cell(2, 1)
        cell_to_populate6 = table.cell(2, 2)

        cell_to_populate7 = table.cell(3, 0)
        cell_to_populate8 = table.cell(3, 2)

        cell_to_populate9 = table.cell(4, 0)
        cell_to_populate10 = table.cell(4, 2)

        cell_to_populate11 = table.cell(5, 0)
        cell_to_populate12 = table.cell(5, 2)

        cell_to_populate13 = table.cell(6, 0)
        cell_to_populate14 = table.cell(6, 2)

        cell_to_populate15 = table.cell(7, 0)
        cell_to_populate16 = table.cell(7, 2)

        cell_to_populate17 = table.cell(8, 0)
        cell_to_populate18 = table.cell(8, 2)

        tennet = table.cell(9, 0)
        consortium = table.cell(9, 2)

        project_title = table.cell(1, 2)
        
        def get_cell_value(index,name):
            value = self.df.loc[index,name]
            return "N/A" if pd.isna(value) else str(value)


        for i, row in self.df.iterrows():
            if originator == row["Originator's Document ID (may further modifications required)"]:

                cell_value1 = get_cell_value(i, "Document Title")
                cell_value2 = get_cell_value(i, "TenneT Document ID")
                cell_value3 = get_cell_value(i, "TenneT Revision")

                cell_value4 = get_cell_value(i, "Originator's Document ID (may further modifications required)")
                cell_value5 = get_cell_value(i, "Originator's Revision")
                cell_value6 = get_cell_value(i, "Asset Document Reference")

                cell_value9 = get_cell_value(i, "Purpose of Submission")
                cell_value10 = get_cell_value(i, "WBS Code")

                cell_value11 = get_cell_value(i, "Purpose of Issue")
                cell_value12 = get_cell_value(i, "WBS Name")

                cell_value13 = get_cell_value(i, "Book")
                cell_value14 = get_cell_value(i, "DCC#")

                cell_value15 = get_cell_value(i, "Chapter")
                cell_value16 = get_cell_value(i, "Document Kind")

                cell_value17 = get_cell_value(i, "Subchapter")
                cell_value18 = get_cell_value(i, "Security Level")

                tennet_beeld = "needed_files\Tennet.jpg"
                consortium_beeld = "needed_files\Consortium.png"
                combined_beeld = "needed_files\Tennet.jpg"

                #header talbe
                header = self.doc.sections[1].header
                table_1 = header.tables[0]
                header_cell = table_1.cell(1, 1)

                #referenced docs table
                table_page_2 = self.doc.tables[3]

                # Populate the cell in the Word document

                cell_to_populate1.text = "Document Title - "
                cell_to_populate1.add_paragraph().add_run(str(cell_value1)).font.color.rgb = RGBColor(0, 112,
                                                                                                    192)  # Light blue color
                cell_to_populate2.text = str(f"TenneT Document ID - ")
                cell_to_populate2.add_paragraph().add_run(str(cell_value2)).font.color.rgb = RGBColor(0, 112,
                                                                                                    192)  # Light blue color
                cell_to_populate3.text = str(f"TenneT revision - ")
                cell_to_populate3.add_paragraph().add_run(str(cell_value3)).font.color.rgb = RGBColor(0, 112,
                                                                                                    192)  # Light blue color

                cell_to_populate4.text = str(f"Originator´s Document ID - ")
                cell_to_populate4.add_paragraph().add_run(str(cell_value4)).font.color.rgb = RGBColor(0, 112,
                                                                                                    192)  # Light blue color
                cell_to_populate5.text = str(f"Originator´s Revision - ")
                cell_to_populate5.add_paragraph().add_run(str(cell_value5)).font.color.rgb = RGBColor(0, 112,
                                                                                                    192)  # Light blue color
                cell_to_populate6.text = str(f"Asset Document Reference - ")
                cell_to_populate6.add_paragraph().add_run(str(cell_value6)).font.color.rgb = RGBColor(0, 112,
                                                                                                    192)  # Light blue color

                cell_to_populate7.text = str(f"Contractor - ")
                cell_to_populate7.add_paragraph().add_run(str("GSC")).font.color.rgb = RGBColor(0, 112,
                                                                                                    192)  # Light blue color
                cell_to_populate8.text = str(f"Contract Number - ")
                cell_to_populate8.add_paragraph().add_run(str("GSC-2GW")).font.color.rgb = RGBColor(0, 112,
                                                                                            192)  # Light blue color

                cell_to_populate9.text = str(f"Purpose of Submission - ")
                cell_to_populate9.add_paragraph().add_run(str(cell_value9)).font.color.rgb = RGBColor(0, 112,
                                                                                                    192)  # Light blue color
                cell_to_populate10.text = str(f"WBS Code - ")
                cell_to_populate10.add_paragraph().add_run(str(cell_value10)).font.color.rgb = RGBColor(0, 112,
                                                                                                    192)  # Light blue color

                cell_to_populate11.text = str(f"Purpose of Issue - ")
                cell_to_populate11.add_paragraph().add_run(str(cell_value11)).font.color.rgb = RGBColor(0, 112,
                                                                                                        192)  # Light blue color
                cell_to_populate12.text = str(f"WBS Name - ")
                cell_to_populate12.add_paragraph().add_run(str(cell_value12)).font.color.rgb = RGBColor(0, 112,
                                                                                                        192)  # Light blue color

                cell_to_populate13.text = str(f"Book - ")
                cell_to_populate13.add_paragraph().add_run(str(cell_value13)).font.color.rgb = RGBColor(0, 112,
                                                                                                        192)  # Light blue color
                cell_to_populate14.text = str(f"DCC# - ")
                cell_to_populate14.add_paragraph().add_run(str(cell_value14)).font.color.rgb = RGBColor(0, 112,
                                                                                                        192)  # Light blue color

                cell_to_populate15.text = str(f"Chapter - ")
                cell_to_populate15.add_paragraph().add_run(str(cell_value15)).font.color.rgb = RGBColor(0, 112,
                                                                                                        192)  # Light blue color
                cell_to_populate16.text = str(f"Chapter - ")
                cell_to_populate16.add_paragraph().add_run(str(cell_value16)).font.color.rgb = RGBColor(0, 112,
                                                                                                        192)  # Light blue color

                cell_to_populate17.text = str(f"Subchapter - ")
                cell_to_populate17.add_paragraph().add_run(str(cell_value17)).font.color.rgb = RGBColor(0, 112,
                                                                                                        192)  # Light blue color
                cell_to_populate18.text = str(f"Security Level - ")
                cell_to_populate18.add_paragraph().add_run(str(cell_value18)).font.color.rgb = RGBColor(0, 112,
                                                                                                        192)  # Light blue color
                #put the pictures in the main table
                tennet_paragraph = tennet.paragraphs[0]
                tennet_run = tennet_paragraph.add_run()
                tennet_run.add_picture(tennet_beeld)
                tennet.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                consortium_paragraph = consortium.paragraphs[0]
                consortium_run = consortium_paragraph.add_run()
                consortium_run.add_picture(consortium_beeld)
                consortium.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

                #header_cell.text = str(f"TenneT Document ID - \n{cell_value2}")
                project_title.text = str(f"Project Name - ")
                project_title.add_paragraph().add_run(str("IJmuiden Ver Beta")).font.color.rgb = RGBColor(0, 112,
                                                                                                    192)  # Light blue color
                
                header_cell.text = str(f"TenneT Document ID - \n{cell_value2}")


                self.export_code_originator = cell_value4

            else:
                pass
        return self.doc, self.export_code_originator
    
    def table_tennet(self, code_tennet):

        """
        This method is used to construct the main table of the coversheet for documents with only a TenneT Document Code

        params: 
        code_tenent: The TenneT Document Code based on the document's titles
        """
        # Get the table and cell to populate
        table = self.doc.tables[0]
        cell_to_populate1 = table.cell(0, 1)
        cell_to_populate2 = table.cell(1, 0)
        cell_to_populate3 = table.cell(1, 1)

        cell_to_populate4 = table.cell(2, 0)
        cell_to_populate5 = table.cell(2, 1)
        cell_to_populate6 = table.cell(2, 2)

        cell_to_populate7 = table.cell(3, 0)
        cell_to_populate8 = table.cell(3, 2)

        cell_to_populate9 = table.cell(4, 0)
        cell_to_populate10 = table.cell(4, 2)

        cell_to_populate11 = table.cell(5, 0)
        cell_to_populate12 = table.cell(5, 2)

        cell_to_populate13 = table.cell(6, 0)
        cell_to_populate14 = table.cell(6, 2)

        cell_to_populate15 = table.cell(7, 0)
        cell_to_populate16 = table.cell(7, 2)

        cell_to_populate17 = table.cell(8, 0)
        cell_to_populate18 = table.cell(8, 2)

        tennet = table.cell(9, 0)
        consortium = table.cell(9, 2)

        project_title = table.cell(1, 2)
        
        def get_cell_value(index,name):
            value = self.df.loc[index,name]
            return "N/A" if pd.isna(value) else str(value)


        for i, row in self.df.iterrows():
            if code_tennet == row["TenneT Document ID"]:

                cell_value1 = get_cell_value(i, "Document Title")
                cell_value2 = get_cell_value(i, "TenneT Document ID")
                cell_value3 = get_cell_value(i, "TenneT Revision")

                cell_value4 = get_cell_value(i, "Originator's Document ID (may further modifications required)")
                cell_value5 = get_cell_value(i, "Originator's Revision")
                cell_value6 = get_cell_value(i, "Asset Document Reference")

                cell_value9 = get_cell_value(i, "Purpose of Submission")
                cell_value10 = get_cell_value(i, "WBS Code")

                cell_value11 = get_cell_value(i, "Purpose of Issue")
                cell_value12 = get_cell_value(i, "WBS Name")

                cell_value13 = get_cell_value(i, "Book")
                cell_value14 = get_cell_value(i, "DCC#")

                cell_value15 = get_cell_value(i, "Chapter")
                cell_value16 = get_cell_value(i, "Document Kind")

                cell_value17 = get_cell_value(i, "Subchapter")
                cell_value18 = get_cell_value(i, "Security Level")

                tennet_beeld = "needed_files\Tennet.jpg"
                consortium_beeld = "needed_files\Consortium.png"
                combined_beeld = "needed_files\Tennet.jpg"

                #header talbe
                header = self.doc.sections[1].header
                table_1 = header.tables[0]
                header_cell = table_1.cell(1, 1)

                #referenced docs table
                table_page_2 = self.doc.tables[3]

                # Populate the cell in the Word document

                cell_to_populate1.text = "Document Title - "
                cell_to_populate1.add_paragraph().add_run(str(cell_value1)).font.color.rgb = RGBColor(0, 112,
                                                                                                    192)  # Light blue color
                cell_to_populate2.text = str(f"TenneT Document ID - ")
                cell_to_populate2.add_paragraph().add_run(str(cell_value2)).font.color.rgb = RGBColor(0, 112,
                                                                                                    192)  # Light blue color
                cell_to_populate3.text = str(f"TenneT revision - ")
                cell_to_populate3.add_paragraph().add_run(str(cell_value3)).font.color.rgb = RGBColor(0, 112,
                                                                                                    192)  # Light blue color

                cell_to_populate4.text = str(f"Originator´s Document ID - ")
                cell_to_populate4.add_paragraph().add_run(str(cell_value4)).font.color.rgb = RGBColor(0, 112,
                                                                                                    192)  # Light blue color
                cell_to_populate5.text = str(f"Originator´s Revision - ")
                cell_to_populate5.add_paragraph().add_run(str(cell_value5)).font.color.rgb = RGBColor(0, 112,
                                                                                                    192)  # Light blue color
                cell_to_populate6.text = str(f"Asset Document Reference - ")
                cell_to_populate6.add_paragraph().add_run(str(cell_value6)).font.color.rgb = RGBColor(0, 112,
                                                                                                    192)  # Light blue color

                cell_to_populate7.text = str(f"Contractor - ")
                cell_to_populate7.add_paragraph().add_run(str("GSC")).font.color.rgb = RGBColor(0, 112,
                                                                                                    192)  # Light blue color
                cell_to_populate8.text = str(f"Contract Number - ")
                cell_to_populate8.add_paragraph().add_run(str("GSC-2GW")).font.color.rgb = RGBColor(0, 112,
                                                                                            192)  # Light blue color

                cell_to_populate9.text = str(f"Purpose of Submission - ")
                cell_to_populate9.add_paragraph().add_run(str(cell_value9)).font.color.rgb = RGBColor(0, 112,
                                                                                                    192)  # Light blue color
                cell_to_populate10.text = str(f"WBS Code - ")
                cell_to_populate10.add_paragraph().add_run(str(cell_value10)).font.color.rgb = RGBColor(0, 112,
                                                                                                    192)  # Light blue color

                cell_to_populate11.text = str(f"Purpose of Issue - ")
                cell_to_populate11.add_paragraph().add_run(str(cell_value11)).font.color.rgb = RGBColor(0, 112,
                                                                                                        192)  # Light blue color
                cell_to_populate12.text = str(f"WBS Name - ")
                cell_to_populate12.add_paragraph().add_run(str(cell_value12)).font.color.rgb = RGBColor(0, 112,
                                                                                                        192)  # Light blue color

                cell_to_populate13.text = str(f"Book - ")
                cell_to_populate13.add_paragraph().add_run(str(cell_value13)).font.color.rgb = RGBColor(0, 112,
                                                                                                        192)  # Light blue color
                cell_to_populate14.text = str(f"DCC# - ")
                cell_to_populate14.add_paragraph().add_run(str(cell_value14)).font.color.rgb = RGBColor(0, 112,
                                                                                                        192)  # Light blue color

                cell_to_populate15.text = str(f"Chapter - ")
                cell_to_populate15.add_paragraph().add_run(str(cell_value15)).font.color.rgb = RGBColor(0, 112,
                                                                                                        192)  # Light blue color
                cell_to_populate16.text = str(f"Chapter - ")
                cell_to_populate16.add_paragraph().add_run(str(cell_value16)).font.color.rgb = RGBColor(0, 112,
                                                                                                        192)  # Light blue color

                cell_to_populate17.text = str(f"Subchapter - ")
                cell_to_populate17.add_paragraph().add_run(str(cell_value17)).font.color.rgb = RGBColor(0, 112,
                                                                                                        192)  # Light blue color
                cell_to_populate18.text = str(f"Security Level - ")
                cell_to_populate18.add_paragraph().add_run(str(cell_value18)).font.color.rgb = RGBColor(0, 112,
                                                                                                        192)  # Light blue color
                #put the pictures in the main table
                tennet_paragraph = tennet.paragraphs[0]
                tennet_run = tennet_paragraph.add_run()
                tennet_run.add_picture(tennet_beeld)
                tennet.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                consortium_paragraph = consortium.paragraphs[0]
                consortium_run = consortium_paragraph.add_run()
                consortium_run.add_picture(consortium_beeld)
                consortium.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

                #header_cell.text = str(f"TenneT Document ID - \n{cell_value2}")
                project_title.text = str(f"Project Name - ")
                project_title.add_paragraph().add_run(str("IJmuiden Ver Beta")).font.color.rgb = RGBColor(0, 112,
                                                                                                    192)  # Light blue color
                
                header_cell.text = str(f"TenneT Document ID - \n{cell_value2}")


                self.export_code_tennet = cell_value2

            else:
                pass
        return self.doc, self.export_code_tennet
        
    def download_originator(self, selected_directory):
        """
        Download function in case the document has only an originato's title

        param:
        selected directory: The location the user wished to save the document
        """
        file_path = f"{selected_directory}/Coversheet- {self.export_code_originator}.docx"
        self.doc.save(file_path)


    def download_tennet(self, selected_directory):
        """
        Download function in case the document has only a TenneT Document Code

        param:
        selected directory: The location the user wished to save the document
        """
        file_path = f"{selected_directory}/Coversheet- {self.export_code_tennet}.docx"
        self.doc.save(file_path)
